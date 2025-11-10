import os
import json
import docx
import PyPDF2
import pandas as pd
from datetime import datetime
from fireworks.client import Fireworks
from openpyxl import Workbook
from openpyxl.styles import PatternFill, Font, Alignment
from sentence_transformers import SentenceTransformer
from qdrant_client import QdrantClient
from qdrant_client.http import models
import numpy as np
import re
import src.config as config

print("✅ Production CV Ranking System v2.0")

# ----------------------------
# Fireworks Setup with Error Handling
# ----------------------------
try:
    fw = Fireworks(api_key=config.FIREWORKS_API_KEY)
    LLM_MODEL = "accounts/fireworks/models/llama-v3p3-70b-instruct"
    EMBED_MODEL = "sentence-transformers/multi-qa-MiniLM-L6-cos-v1"
    sbert_model = SentenceTransformer(EMBED_MODEL)
    qdrant = QdrantClient("http://localhost:6333")
    COLLECTION_NAME = "cv_ranking"
    print("✅ AI services initialized successfully")
except Exception as e:
    print(f"⚠️ Warning: AI initialization error: {e}")
    fw = None
    sbert_model = None
    qdrant = None

# ----------------------------
# CV Validation Functions
# ----------------------------
def is_arabic_text(text):
    """Check if text contains significant Arabic characters"""
    if not text:
        return False
    arabic_chars = len(re.findall(r'[\u0600-\u06FF]', text))
    total_chars = len(re.findall(r'[a-zA-Z\u0600-\u06FF]', text))
    
    if total_chars > 0 and (arabic_chars / total_chars) > 0.3:
        return True
    return False

def has_sufficient_english_content(text):
    """Check if text has sufficient English content"""
    if not text:
        return False
    
    english_words = re.findall(r'\b[a-zA-Z]+\b', text)
    return len(english_words) >= 30  # Lowered threshold for production

def is_valid_cv_content(text):
    """
    Production CV validation with balanced rules.
    Returns: (is_valid: bool, reason: str)
    """
    if not text or len(text.strip()) < 50:
        return False, "CV is empty or too short (less than 50 characters)"
    
    # Check for sufficient English content
    english_words = re.findall(r'\b[a-zA-Z]+\b', text)
    if len(english_words) < 30:
        return False, "CV lacks sufficient English content (minimum 30 words required)"
    
    # Check for common CV indicators (at least 1 should be present)
    cv_indicators = [
        r'\b(experience|education|skills|qualifications|employment|work history)\b',
        r'\b(bachelor|master|phd|degree|university|college|diploma|certification)\b',
        r'\b(company|organization|firm|startup|enterprise|corporation)\b',
        r'\b(email|phone|address|contact|mobile)\b',
        r'\b(january|february|march|april|may|june|july|august|september|october|november|december|\d{4})\b'
    ]
    
    indicators_found = sum(1 for pattern in cv_indicators if re.search(pattern, text.lower()))
    
    if indicators_found < 1:
        return False, "CV does not contain expected sections (education, experience, contact info, etc.)"
    
    # Check for excessive special characters (up to 50% allowed)
    special_char_ratio = len(re.findall(r'[^a-zA-Z0-9\s\.,;:\-\(\)\[\]\/]', text)) / max(len(text), 1)
    if special_char_ratio > 0.5:
        return False, "CV contains excessive special characters - may be corrupted"
    
    return True, "Valid"

def validate_cv_file(cv_data):
    """
    Main CV validation function.
    Returns: (is_valid: bool, reason: str, cleaned_text: str)
    """
    text = cv_data.get("text", "")
    filename = cv_data.get("filename", "unknown")
    
    # Basic validation
    if not text or text.strip() == "":
        return False, "Failed to extract text - file may be image-based, password-protected, or corrupted", ""
    
    # Content validation
    is_valid, reason = is_valid_cv_content(text)
    
    if not is_valid:
        print(f"   ⚠️ Invalid CV: {reason}")
        return False, reason, ""
    
    return True, "Valid", text

# ----------------------------
# CV Parsing Functions with Error Handling
# ----------------------------
def extract_text_from_pdf(file_path):
    """Extract text from PDF with comprehensive error handling"""
    try:
        if not os.path.exists(file_path):
            print(f"❌ File not found: {file_path}")
            return ""
        
        text = ""
        with open(file_path, "rb") as f:
            try:
                reader = PyPDF2.PdfReader(f)
                
                # Check if PDF is encrypted
                if reader.is_encrypted:
                    print(f"⚠️ Password-protected PDF: {file_path}")
                    return ""
                
                # Check if PDF has pages
                if len(reader.pages) == 0:
                    print(f"⚠️ Empty PDF: {file_path}")
                    return ""
                
                # Extract text from all pages
                for page in reader.pages:
                    try:
                        page_text = page.extract_text() or ""
                        text += page_text
                    except:
                        continue
                        
            except PyPDF2.errors.PdfReadError:
                print(f"❌ Corrupted PDF: {file_path}")
                return ""
            except Exception as e:
                print(f"❌ PDF error: {str(e)[:50]}")
                return ""
        
        return text.strip()
        
    except Exception as e:
        print(f"❌ Unexpected PDF error: {str(e)[:50]}")
        return ""

def extract_text_from_docx(file_path):
    """Extract text from DOCX with error handling"""
    try:
        if not os.path.exists(file_path):
            print(f"❌ File not found: {file_path}")
            return ""
        
        doc = docx.Document(file_path)
        
        # Extract from paragraphs
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        
        # Extract from tables
        tables_text = []
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        tables_text.append(cell.text)
        
        text = "\n".join(paragraphs + tables_text)
        return text.strip()
        
    except Exception as e:
        print(f"❌ DOCX error {file_path}: {str(e)[:50]}")
        return ""

def extract_text_from_txt(file_path):
    """Extract text from TXT with encoding fallback"""
    encodings = ['utf-8', 'latin-1', 'windows-1252', 'iso-8859-1']
    
    for encoding in encodings:
        try:
            with open(file_path, 'r', encoding=encoding) as f:
                return f.read().strip()
        except UnicodeDecodeError:
            continue
        except Exception as e:
            print(f"❌ TXT error: {str(e)[:50]}")
            return ""
    
    print(f"❌ Cannot decode TXT file: {file_path}")
    return ""

def extract_text(file_path):
    """Extract text from PDF, DOCX, DOC, or TXT files"""
    if not file_path or not os.path.exists(file_path):
        return ""
    
    ext = os.path.splitext(file_path)[1].lower()
    
    if ext == ".pdf":
        return extract_text_from_pdf(file_path)
    elif ext in [".docx", ".doc"]:
        return extract_text_from_docx(file_path)
    elif ext == ".txt":
        return extract_text_from_txt(file_path)
    else:
        print(f"⚠️ Unsupported file type: {ext}")
        return ""

# ----------------------------
# Job Description Loading
# ----------------------------
def load_job_description(job_description_input):
    """Load job description from file or text"""
    try:
        # If it's a file path
        if isinstance(job_description_input, str) and os.path.exists(job_description_input):
            print(f"📄 Loading job description from: {job_description_input}")
            text = extract_text(job_description_input)
            if text:
                print(f"✅ Job description loaded ({len(text)} characters)")
                return text
            else:
                print("⚠️ Could not extract job description text")
                return ""
        
        # If it's direct text
        elif isinstance(job_description_input, str) and job_description_input.strip():
            print(f"📝 Using job description text ({len(job_description_input)} characters)")
            return job_description_input.strip()
        
        print("⚠️ Empty job description")
        return ""
        
    except Exception as e:
        print(f"❌ Job description error: {e}")
        return ""

def load_cvs_from_dataframe(df):
    """Load CVs from DataFrame"""
    candidates = []
    
    if isinstance(df, list):
        df = pd.DataFrame(df)
    
    for idx, row in df.iterrows():
        try:
            cv_path = row.get("local_cv_path")
            cv_link = row.get("CV", "")
            name = (row.get("Full Name") or row.get("full name") or 
                   row.get("fullname") or f"Candidate_{idx+1}")
            
            text = ""
            if cv_path and os.path.exists(cv_path):
                text = extract_text(cv_path)
            
            candidates.append({
                "filename": os.path.basename(cv_path) if cv_path else f"candidate_{idx+1}.pdf",
                "text": text,
                "name": str(name).strip(),
                "cv_link": cv_link
            })
            
        except Exception as e:
            print(f"❌ Error loading CV at row {idx}: {str(e)[:50]}")
            continue
    
    print(f"📋 Loaded {len(candidates)} CVs")
    return candidates

# ----------------------------
# AI-Powered Ranking
# ----------------------------
def rank_with_gemini(cvs, job_description, api_key=None, batch_size=1):
    """
    Rank CVs using Fireworks AI with Llama 3.3 70B
    Universal for all specialties with production error handling
    """
    jd_text = load_job_description(job_description)
    
    if not jd_text:
        print("⚠️ Empty job description - rankings may be less accurate")
    
    results = []
    
    if not api_key or not fw:
        print("❌ AI service unavailable")
        for c in cvs:
            results.append({
                "filename": c["filename"],
                "name": c["name"],
                "score": 0,
                "reasoning": "❌ AI service unavailable",
                "cv_link": c.get("cv_link", "")
            })
        return results
    
    # Track statistics
    total = len(cvs)
    valid = 0
    invalid = 0
    errors = 0
    
    for i, cv in enumerate(cvs):
        try:
            print(f"🔥 Analyzing CV {i+1}/{total}: {cv['name']}")
            
            # VALIDATE CV
            is_valid, validation_reason, cleaned_text = validate_cv_file(cv)
            
            if not is_valid:
                invalid += 1
                results.append({
                    "filename": cv["filename"],
                    "name": cv["name"],
                    "score": 0,
                    "reasoning": f"❌ INVALID: {validation_reason}",
                    "cv_link": cv.get("cv_link", "")
                })
                print(f"   ❌ Invalid: {validation_reason}")
                continue
            
            valid += 1
            print(f"   ✅ Valid - analyzing...")
            
            # AI Analysis
            prompt = f"""You are an expert recruiter. Analyze this CV against the job requirements.

JOB REQUIREMENTS:
{jd_text[:1500]}

CANDIDATE CV:
Name: {cv['name']}
{cv['text'][:2000]}

EVALUATION CRITERIA:
1. Qualifications & Education (degrees, certifications, licenses)
2. Relevant Experience (years, industry match, progression)
3. Technical/Professional Skills (tools, methodologies, competencies)
4. Achievements & Impact (results, projects, awards)
5. Soft Skills & Fit (communication, teamwork, problem-solving)

SCORING SCALE:
- 90-100: Excellent match (perfect fit, extensive experience)
- 75-89: Strong match (right qualifications, good experience)
- 60-74: Good match (relevant background, some experience)
- 40-59: Moderate match (general background, limited match)
- 20-39: Weak match (different field, insufficient experience)
- 0-19: Poor match (unrelated background, major gaps)

Respond ONLY with JSON:
{{"score": 85, "reasoning": "Brief explanation of match quality"}}"""

            response = fw.chat.completions.create(
                model=LLM_MODEL,
                messages=[
                    {"role": "system", "content": "You are an expert recruiter. Respond only with valid JSON."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.2,
                max_tokens=400
            )
            
            raw_text = response.choices[0].message.content.strip()
            
            # Parse JSON
            parsed = None
            try:
                clean_text = re.sub(r'```(?:json)?\s*', '', raw_text)
                clean_text = re.sub(r'```', '', clean_text).strip()
                parsed = json.loads(clean_text)
            except json.JSONDecodeError:
                json_match = re.search(r'\{[^{}]*"score"\s*:\s*\d+[^{}]*"reasoning"\s*:[^{}]*\}', raw_text, re.DOTALL)
                if json_match:
                    try:
                        parsed = json.loads(json_match.group(0))
                    except:
                        pass
            
            if parsed and "score" in parsed and "reasoning" in parsed:
                score = max(0, min(100, int(parsed["score"])))
                reasoning = parsed["reasoning"].strip()
                
                results.append({
                    "filename": cv["filename"],
                    "name": cv["name"],
                    "score": score,
                    "reasoning": reasoning,
                    "cv_link": cv.get("cv_link", "")
                })
                print(f"   ✅ Score: {score}/100")
            else:
                errors += 1
                print(f"   ⚠️ Parse failed")
                results.append({
                    "filename": cv["filename"],
                    "name": cv["name"],
                    "score": 0,
                    "reasoning": "❌ Analysis failed - review manually",
                    "cv_link": cv.get("cv_link", "")
                })
                
        except Exception as e:
            errors += 1
            print(f"   ❌ Error: {str(e)[:80]}")
            results.append({
                "filename": cv["filename"],
                "name": cv["name"],
                "score": 0,
                "reasoning": f"❌ Error: {str(e)[:80]}",
                "cv_link": cv.get("cv_link", "")
            })
    
    print(f"\n📊 Summary: {total} total | {valid} valid | {invalid} invalid | {errors} errors")
    return results

# ----------------------------
# Embedding-based Ranking
# ----------------------------
def rank_with_embeddings(cvs, job_description, top_k=5):
    """Rank CVs using semantic embeddings"""
    jd_text = load_job_description(job_description)
    
    if not jd_text:
        print("⚠️ Empty job description")
        jd_text = "General position"
    
    valid_cvs = []
    invalid_results = []
    
    for cv in cvs:
        is_valid, validation_reason, cleaned_text = validate_cv_file(cv)
        if is_valid:
            valid_cvs.append(cv)
        else:
            invalid_results.append({
                "filename": cv["filename"],
                "name": cv["name"],
                "score": 0,
                "reasoning": f"INVALID: {validation_reason}",
                "cv_link": cv.get("cv_link", "")
            })
    
    if len(valid_cvs) == 0:
        print("⚠️ No valid CVs")
        return invalid_results
    
    if not qdrant or not sbert_model:
        print("❌ Embedding service unavailable")
        return invalid_results
    
    try:
        if not qdrant.collection_exists(COLLECTION_NAME):
            qdrant.recreate_collection(
                collection_name=COLLECTION_NAME,
                vectors_config=models.VectorParams(size=384, distance=models.Distance.COSINE),
            )
        
        qdrant.delete_collection(COLLECTION_NAME)
        qdrant.recreate_collection(
            collection_name=COLLECTION_NAME,
            vectors_config=models.VectorParams(size=384, distance=models.Distance.COSINE),
        )
        
        vectors = []
        for idx, c in enumerate(valid_cvs):
            emb = sbert_model.encode(c["text"] or "")
            vectors.append(models.PointStruct(
                id=idx,
                vector=emb.tolist(),
                payload={"filename": c["filename"], "name": c["name"], "cv_link": c["cv_link"], "text": c["text"]}
            ))
        
        qdrant.upsert(collection_name=COLLECTION_NAME, points=vectors)
        
        job_emb = sbert_model.encode(jd_text).tolist()
        search_results = qdrant.search(
            collection_name=COLLECTION_NAME,
            query_vector=job_emb,
            limit=min(top_k, len(valid_cvs))
        )
        
        results = []
        for r in search_results:
            payload = r.payload
            score = round(r.score * 100, 2)
            results.append({
                "filename": payload.get("filename", "Unknown"),
                "name": payload.get("name", "Unknown"),
                "score": score,
                "reasoning": f"Semantic similarity: {score}%",
                "cv_link": payload.get("cv_link", "")
            })
        
        return results + invalid_results
        
    except Exception as e:
        print(f"❌ Embedding error: {e}")
        return invalid_results

# ----------------------------
# Enhanced Excel Export
# ----------------------------
def save_results_to_excel(results_list, job_description=None, output_dir=".", output_path=None):
    """Save ranking results to Excel with professional formatting"""
    
    df_out = pd.DataFrame(results_list)
    
    required_cols = ["name", "score", "status", "reasoning", "cv_link"]
    for col in required_cols:
        if col not in df_out.columns:
            df_out[col] = ""
    
    df_out["status"] = df_out["score"].apply(lambda x: "Match" if float(x) >= 60 else "No Match")
    df_out = df_out[required_cols].sort_values(by="score", ascending=False)
    
    wb = Workbook()
    ws = wb.active
    ws.title = "Ranked Candidates"
    
    headers = list(df_out.columns)
    ws.append(headers)
    
    # Header formatting
    header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
    header_font = Font(bold=True, color="FFFFFF", size=12)
    for cell in ws[1]:
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center")
    
    # Data rows
    for _, row in df_out.iterrows():
        ws.append([row.get(h, "") for h in headers])
    
    # Color coding
    for i in range(2, ws.max_row + 1):
        status = ws[f"C{i}"].value
        try:
            score = float(ws[f"B{i}"].value)
        except:
            score = 0
        
        if status == "Match":
            fill = PatternFill(start_color="C6EFCE", fill_type="solid")
        elif score >= 40:
            fill = PatternFill(start_color="FFEB9C", fill_type="solid")
        else:
            fill = PatternFill(start_color="F2DCDB", fill_type="solid")
        
        for cell in ws[i]:
            cell.fill = fill
        
        # Clickable links
        cv_link = ws[f"E{i}"].value
        if cv_link and str(cv_link).startswith("http"):
            ws[f"E{i}"].hyperlink = cv_link
            ws[f"E{i}"].style = "Hyperlink"
    
    # Column widths
    ws.column_dimensions['A'].width = 25
    ws.column_dimensions['B'].width = 10
    ws.column_dimensions['C'].width = 15
    ws.column_dimensions['D'].width = 60
    ws.column_dimensions['E'].width = 40
    
    # Summary Sheet
    ws_summary = wb.create_sheet("Summary")
    
    total = len(df_out)
    valid = len(df_out[df_out["score"] > 0])
    invalid = len(df_out[df_out["score"] == 0])
    avg_score = df_out[df_out["score"] > 0]["score"].astype(float).mean() if valid > 0 else 0
    top = df_out.iloc[0] if not df_out.empty else None
    
    ws_summary["A1"], ws_summary["B1"] = "Total CVs", total
    ws_summary["A2"], ws_summary["B2"] = "Valid CVs", valid
    ws_summary["A3"], ws_summary["B3"] = "Invalid CVs", invalid
    ws_summary["A4"], ws_summary["B4"] = "Average Score", round(avg_score, 1)
    
    if top is not None:
        ws_summary["A6"], ws_summary["B6"] = "Top Candidate", top.get("name", "Unknown")
        ws_summary["A7"], ws_summary["B7"] = "Top Score", top.get("score", 0)
        ws_summary["A8"], ws_summary["B8"] = "CV Link", top.get("cv_link", "")
    
    # Generate path
    if output_path:
        final_path = output_path
    else:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        final_path = os.path.join(output_dir, f"CV_Rankings_{timestamp}.xlsx")
    
    wb.save(final_path)
    print(f"✅ Results saved: {final_path}")
    print(f"📊 {valid} valid | {invalid} invalid | Avg: {avg_score:.1f}/100")
    
    return df_out, final_path